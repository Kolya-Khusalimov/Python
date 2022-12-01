import docx
import os
import sys

#23.1
def copy_style(folder, basefilename):
    full_name = os.path.join(folder, basefilename)
    document = Document(full_name)
    paragraph = document.paragraphs[0]
    result = {}
    items = ()

    for i in items:
        if it in paragraph.style:
            result[it] = paragraph.style[it]
    return result

def apply_style(copied_style, folder, basefilename):
    os.chdir(folder)

    for path, files, dirs in os.walk():
        for doc_file in files:
            if (doc_file[-5:] != '.docx'):
                continue

            if doc_file == basefilename:
                continue

            document = Document(filename)
            for paragraph in document.paragraphs:
                for k, v in copied_style.items():
                    if k in paragraph.style:
                        paragraph.style[k] = v

            document.save(os.path.join())


if __name__ == '__main__':
    if len(sys.argv) == 1:
        basefilename = input("template.docx file name: ")
        folder = input("folder name")
    elif len(sys.argv) == 2:
        basefilename = sys.argv(1)
        folder = "."

    else:
        basefilename = sys.argv(1)
        folder = sys.argv[2]


    copied_style = copy_style(folder, basefilename)
    apply_style(copied_style, folder, basefilename)



#23.2
import re

def change_dates(text):
    print(match)
    date = match.group()
    if "/" in date:
        y, m, d = date.split("/")
    else:
        d, m, y = date.split(".")
    while len(y) != 4:
        y = "0" + y
    if len(m) != 2:
        m = "0" + m
    if len(d) != 2:
        d = "0" + d
    date = ".".join((d, m, y))
    return date



def format_dates(filename):
    document = Document(full_name)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:  #переглядаємо частину параграфа
            text = run.text  #
            run.text = change_dates(text) #

    document.save(full_name)

if __name__ == '__main__':
    if len(sys.argv) == 1:
            filename = input(".docx file name: ")
    else:
        filename = sys.argv[1]
    format_dates(filename)


#23.9
from openpyxl import *

def output_cells(folder, text):
    os.chdir(folder)
    for path, files, dirs in os.walk():
        for item in files:
            if item.endwith(".xlsx"):
                if not item.endwith(".xlsx"):
                    continue
                path = os.path.join(path, item)
                try:
                    wb = load_workbook(full_name)
                    lst = wb.get_sheet_names()
                    for sheet in lst:
                        #print(f"sheet{sheet}")
                        #ws = wb[sheet]

                        for row in ws.iter_rows():
                            for cell in row:
                                cell_text = cell.value
                                if not cell_text:
                                    continue
                                    print(f"t={cell_text}")
                                    if re.match(cell_text, text):
                                       print(f"{fullname}, {sheet}, cell {row}, {column} has text")
                except:
                    print(f"can\'t open {path}")


if __name__ == '__main__':
    if len(sys.argv)>=3:
        folder = sys.argv[1]
        text = sys.argv[2]
    else:
        folder = input("folder: ")
        text = input("to find: ")
